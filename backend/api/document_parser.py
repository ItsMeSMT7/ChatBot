# d:\Sumit\Project\P Square\backend\api\document_parser.py

"""
Smart document parser.
Extracts structured text (headings, paragraphs, page numbers)
from PDF, DOCX, and TXT files.
"""

import os
import re
import logging

logger = logging.getLogger(__name__)


class ParsedSection:
    """One logical block of text from the document."""

    def __init__(self, content, heading=None, page_number=None,
                 section_type="paragraph"):
        self.content = content          # the actual text
        self.heading = heading          # heading this text falls under
        self.page_number = page_number  # which page (PDF only)
        self.section_type = section_type  # "heading" or "paragraph"


class ParsedDocument:
    """Complete parsed output of a document."""

    def __init__(self):
        self.sections = []      # list of ParsedSection
        self.full_text = ""     # entire document as one string
        self.total_pages = 0
        self.title = ""


# ────────────────────────────────────────────────────────────
#  PDF Parsing  (PyMuPDF)
# ────────────────────────────────────────────────────────────
def _parse_pdf(file_path):
    """
    Extract text from PDF page-by-page.
    Detects headings using font-size analysis.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed, using fallback PDF reader")
        return _parse_pdf_fallback(file_path)

    doc = fitz.open(file_path)
    parsed = ParsedDocument()
    parsed.total_pages = len(doc)
    parsed.title = os.path.basename(file_path)

    all_text = []
    current_heading = None

    for page_num in range(len(doc)):
        page = doc[page_num]
        blocks = page.get_text("dict", sort=True).get("blocks", [])

        # Find the most common font size on this page (= body text size)
        font_sizes = []
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    font_sizes.append(span["size"])

        if not font_sizes:
            continue

        # Body text = most common size; anything bigger = heading
        font_sizes.sort()
        body_size = font_sizes[len(font_sizes) // 2]
        heading_threshold = body_size + 1.5

        for block in blocks:
            if "lines" not in block:
                continue

            block_text = ""
            is_heading = False

            for line in block["lines"]:
                for span in line["spans"]:
                    block_text += span["text"]
                    if (span["size"] >= heading_threshold
                            or "bold" in span["font"].lower()):
                        is_heading = True

            block_text = block_text.strip()
            if not block_text:
                continue

            # Short bold/large text → heading
            if is_heading and len(block_text) < 200:
                current_heading = block_text
                parsed.sections.append(ParsedSection(
                    content=block_text,
                    heading=block_text,
                    page_number=page_num + 1,
                    section_type="heading",
                ))
            else:
                parsed.sections.append(ParsedSection(
                    content=block_text,
                    heading=current_heading,
                    page_number=page_num + 1,
                    section_type="paragraph",
                ))

            all_text.append(block_text)

    doc.close()
    parsed.full_text = "\n".join(all_text)
    return parsed


def _parse_pdf_fallback(file_path):
    """Fallback PDF reader using pdfplumber."""
    parsed = ParsedDocument()
    parsed.title = os.path.basename(file_path)

    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            parsed.total_pages = len(pdf.pages)
            all_text = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    parsed.sections.append(ParsedSection(
                        content=text.strip(),
                        page_number=i + 1,
                    ))
                    all_text.append(text.strip())
            parsed.full_text = "\n".join(all_text)
    except ImportError:
        # Last resort: PyPDF2 (likely already installed)
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            parsed.total_pages = len(reader.pages)
            all_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    parsed.sections.append(ParsedSection(
                        content=text.strip(),
                        page_number=i + 1,
                    ))
                    all_text.append(text.strip())
            parsed.full_text = "\n".join(all_text)
        except Exception as e:
            logger.error(f"All PDF parsers failed: {e}")

    return parsed


# ────────────────────────────────────────────────────────────
#  DOCX Parsing
# ────────────────────────────────────────────────────────────
def _parse_docx(file_path):
    """Parse Word document, detecting headings by paragraph style."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parsed = ParsedDocument()
    parsed.title = os.path.basename(file_path)

    current_heading = None
    all_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = (para.style.name or "").lower()

        if "heading" in style or "title" in style:
            current_heading = text
            parsed.sections.append(ParsedSection(
                content=text,
                heading=text,
                section_type="heading",
            ))
        else:
            parsed.sections.append(ParsedSection(
                content=text,
                heading=current_heading,
                section_type="paragraph",
            ))
        all_text.append(text)

    parsed.full_text = "\n".join(all_text)
    parsed.total_pages = max(1, len(all_text) // 40)
    return parsed


# ────────────────────────────────────────────────────────────
#  TXT Parsing
# ────────────────────────────────────────────────────────────
def _parse_txt(file_path):
    """
    Parse plain text file.
    Detects headings by patterns: ALL CAPS lines, numbered sections,
    markdown headings.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    parsed = ParsedDocument()
    parsed.title = os.path.basename(file_path)
    parsed.full_text = text
    parsed.total_pages = 1

    heading_patterns = [
        r'^#{1,6}\s+',               # Markdown: ## Heading
        r'^[A-Z][A-Z\s]{4,}$',       # ALL CAPS LINE
        r'^\d+\.\s+[A-Z]',           # 1. Introduction
        r'^(Section|Article|Chapter|PART)\s+\d+',
    ]

    current_heading = None
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        is_heading = any(re.match(p, stripped) for p in heading_patterns)

        if is_heading:
            current_heading = stripped
            parsed.sections.append(ParsedSection(
                content=stripped,
                heading=stripped,
                section_type="heading",
            ))
        else:
            parsed.sections.append(ParsedSection(
                content=stripped,
                heading=current_heading,
                section_type="paragraph",
            ))

    return parsed


# ────────────────────────────────────────────────────────────
#  Main Entry Point
# ────────────────────────────────────────────────────────────
def parse_document(file_path):
    """
    Detect file type and route to the correct parser.
    Returns a ParsedDocument object.
    """
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Parsing: {file_path} (type: {ext})")

    parsers = {
        '.pdf': _parse_pdf,
        '.docx': _parse_docx,
        '.doc': _parse_docx,
        '.txt': _parse_txt,
        '.md': _parse_txt,
    }

    parser_fn = parsers.get(ext)
    if parser_fn is None:
        logger.warning(f"No parser for {ext}, reading as plain text")
        return _parse_txt(file_path)

    try:
        result = parser_fn(file_path)
        logger.info(
            f"Parsed '{result.title}': "
            f"{len(result.sections)} sections, {result.total_pages} pages"
        )
        return result
    except Exception as e:
        logger.error(f"Parser failed: {e}")
        return _parse_txt(file_path)